from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "deliverables"
IMG_DIR = ROOT / "assets"


REPORT_FILE = OUT_DIR / "Automatic_Headlight_Control_Report.docx"
IEEE_FILE = OUT_DIR / "Automatic_Headlight_Control_IEEE_Paper.docx"


TITLE_TEXT = "Automatic Headlight Control System Using Simulink and Stateflow"
PROJECT_NAME = "Automatic Headlight Control"
AUTHOR_PLACEHOLDER = "[Your Name]"
PRN_PLACEHOLDER = "[Your PRN]"
DEPT_PLACEHOLDER = "Department of Electronics and Telecommunication Engineering"
INST_PLACEHOLDER = "[Institute Name]"
COURSE_TEXT = "Model-Based Design using Simulink & Stateflow"


ABSTRACT_TEXT = (
    "This project presents an Automatic Headlight Control System implemented using "
    "Simulink and Stateflow for an automotive application. The model senses ambient "
    "light intensity, vehicle speed, manual override input, and fault status, then "
    "generates the final headlamp command. The complete design is divided into four "
    "subsystems: Input Subsystem, Logic Subsystem, Stateflow Controller, and Output "
    "Subsystem. The Input Subsystem creates four simulation scenarios: day moving, "
    "night moving, tunnel entry, and night stopped. The Logic Subsystem processes raw "
    "signals into boolean status signals using relational operators and hysteresis "
    "logic. Hysteresis is implemented with two thresholds, 300 lux and 500 lux, to "
    "prevent rapid switching of headlights near the decision boundary. The Stateflow "
    "Controller uses four states, OFF, ON, HOLD, and MANUAL, to generate the final "
    "headlamp command. The Output Subsystem displays, plots, and logs the result for "
    "verification. Simulation results confirm correct behavior across all scenarios. "
    "The model demonstrates how model-based design can be used to build and validate "
    "a practical automotive control system."
)


INTRO_PARAS = [
    (
        "Modern vehicles use embedded control systems to improve safety, automation, "
        "and driver convenience. One common example is automatic headlight control, "
        "which turns the headlights ON or OFF according to surrounding brightness and "
        "vehicle condition."
    ),
    (
        "In practical driving, the vehicle may move in daylight, drive at night, enter "
        "a tunnel, stop at a signal in darkness, or require direct manual control from "
        "the driver. An automatic controller reduces dependency on driver attention and "
        "improves response to changing environmental conditions."
    ),
    (
        "This project models an Automatic Headlight Control System in Simulink and "
        "Stateflow. Simulink is used for subsystem-based signal processing and scenario "
        "generation, while Stateflow is used for state-based decision making. The final "
        "output is the signal headlamp_cmd, where 0 represents headlights OFF and 1 "
        "represents headlights ON."
    ),
]


PROBLEM_STATEMENT = (
    "The objective of this project is to design and simulate an automotive headlight "
    "controller that reacts correctly to environmental light and vehicle movement. The "
    "system must identify bright and dark conditions, distinguish between moving and "
    "stopped vehicle states, allow driver override, and generate a correct final "
    "headlamp command for multiple operating scenarios."
)


SYSTEM_ARCHITECTURE_TEXT = (
    "The model follows a modular architecture: Input Subsystem -> Logic Subsystem -> "
    "Stateflow Controller -> Output Subsystem. This organization separates scenario "
    "generation, boolean decision preparation, state-based control, and output "
    "verification."
)


INPUT_SUBSYSTEM_PARAS = [
    (
        "The Input Subsystem creates the operating conditions required for simulation. "
        "A Scenario Constant block selects one of four cases: day moving, night moving, "
        "tunnel entry, and night stopped. Based on the selected scenario, the subsystem "
        "produces AmbientLux and VehicleSpeed signals."
    ),
    (
        "AmbientLux represents environmental brightness. Day condition is modeled with "
        "800 lux, while night condition is modeled with 100 lux. In the tunnel-entry "
        "scenario, the lux value starts at 800, drops by 700 at approximately 10 s, "
        "and returns to 800 at approximately 20 s, simulating tunnel entry and exit."
    ),
    (
        "VehicleSpeed is 60 for moving scenarios and 0 for the stopped-at-night "
        "scenario. ManualOverride is generated from a manual switch and is passed as a "
        "boolean driver command. Fault is currently held at 0 and passed through the "
        "model as an available supervisory input."
    ),
]


LOGIC_SUBSYSTEM_PARAS = [
    (
        "The Logic Subsystem converts raw input values into boolean signals that can be "
        "used directly by the Stateflow Controller. The most important part is the "
        "ambient-light hysteresis logic, which avoids rapid ON/OFF switching near the "
        "decision threshold."
    ),
    (
        "Two relational operators are used: AmbientLux < 300 and AmbientLux > 500. The "
        "first detects dark conditions and the second detects bright conditions. A Unit "
        "Delay stores the previous lux_status value. NOT, AND, and OR blocks are then "
        "connected to form the latch equation lux_status = "
        "(previous_lux_status AND NOT lux_gt_500) OR lux_lt_300."
    ),
    (
        "VehicleSpeed > 0 generates speed_status. ManualOverride is passed through as "
        "override, and Fault is passed through as fault. The final outputs of this "
        "subsystem are lux_status, speed_status, override, and fault."
    ),
]


STATEFLOW_PARAS = [
    (
        "The Stateflow Controller receives lux_status, speed_status, override, and "
        "fault, and generates the final headlamp_cmd output."
    ),
    (
        "The chart contains four states: OFF, ON, HOLD, and MANUAL. OFF is active when "
        "brightness is sufficient and headlights are not needed. ON is active when it "
        "is dark and the vehicle is moving. HOLD is active when it is dark and the "
        "vehicle is stopped, keeping the lamps ON for safety and visibility. MANUAL is "
        "active whenever ManualOverride is asserted and has highest priority."
    ),
    (
        "The decision sequence is straightforward: if override equals 1, the chart goes "
        "to MANUAL. Otherwise, if lux_status equals 0, the chart goes to OFF. If "
        "lux_status equals 1 and speed_status equals 1, the chart goes to ON. If "
        "lux_status equals 1 and speed_status equals 0, the chart goes to HOLD."
    ),
]


OUTPUT_PARAS = [
    (
        "The Output Subsystem receives the final headlamp_cmd signal from Stateflow and "
        "presents it in three ways."
    ),
    (
        "Headlamp_Display shows the current logical value of the command. A value of 0 "
        "means headlights OFF, and a value of 1 means headlights ON."
    ),
    (
        "Headlamp_Scope plots headlamp_cmd over time, making scenario-based transitions "
        "easy to verify visually. Headlamp_ToWorkspace logs the output signal as "
        "out.headlamp_cmd_ts for simulation verification and reporting."
    ),
]


CONCLUSION_TEXT = (
    "The Automatic Headlight Control model successfully demonstrates an automotive "
    "control application using Simulink and Stateflow. The model senses operating "
    "conditions, processes those signals through hysteresis-based logic, makes a "
    "state-based decision, and verifies the output through display, scope, and "
    "workspace logging. The simulation results confirm correct response in daylight, "
    "night driving, tunnel transition, and stopped-at-night conditions."
)


IEEE_INTRO = INTRO_PARAS + [
    (
        "The main contribution of this work is a clear model-based implementation of a "
        "state-driven headlight controller with hysteresis and manual override, tested "
        "through four representative scenarios."
    )
]


IEEE_LIT_REVIEW = [
    (
        "Model-Based Design is widely used in automotive control development because it "
        "supports modular modeling, simulation, early validation, and direct linkage "
        "between control design and implementation [1]."
    ),
    (
        "Stateflow is particularly suitable for supervisory automotive controllers "
        "because it expresses operating modes, transitions, and priority decisions in a "
        "clear finite-state form [2], [3]."
    ),
    (
        "Automatic headlight control systems rely on ambient-light sensing and decision "
        "thresholds to control lamps according to external brightness conditions [4]."
    ),
    (
        "Prior automotive lighting studies also show the importance of robust control "
        "logic for safety-related vehicle lighting behavior and adaptive lighting "
        "functions [5], [6]."
    ),
]


IEEE_METHOD = [
    (
        "The model is built using four connected subsystems: Input Subsystem, Logic "
        "Subsystem, Stateflow Controller, and Output Subsystem. The simulation is "
        "configured as a discrete-time model with fixed step size 0.1 s and stop time "
        "40 s."
    ),
    (
        "The Input Subsystem generates AmbientLux, VehicleSpeed, ManualOverride, and "
        "Fault. Four scenarios are used for testing: day moving, night moving, tunnel "
        "entry, and night stopped."
    ),
    (
        "The Logic Subsystem converts raw signals to boolean supervisory inputs. "
        "AmbientLux < 300 and AmbientLux > 500 form the lower and upper thresholds for "
        "hysteresis. A Unit Delay and logical operators implement the lux-status latch. "
        "VehicleSpeed > 0 creates speed_status, while ManualOverride and Fault are "
        "passed through."
    ),
    (
        "The Stateflow chart implements OFF, ON, HOLD, and MANUAL states. Manual "
        "override has the highest priority. OFF drives headlamp_cmd = 0, while ON, "
        "HOLD, and MANUAL drive headlamp_cmd = 1."
    ),
]


IEEE_RESULTS = [
    (
        "Scenario 1 confirms correct OFF behavior in bright daylight with vehicle "
        "motion. Scenario 2 confirms correct ON behavior in dark conditions while the "
        "vehicle is moving."
    ),
    (
        "Scenario 3 verifies transient response to a tunnel event. The output remains "
        "LOW before approximately 10 s, goes HIGH between approximately 10 s and 20 s "
        "when brightness drops, and returns LOW after approximately 20 s when the "
        "vehicle exits the tunnel."
    ),
    (
        "Scenario 4 verifies that the HOLD state keeps headlights ON when the vehicle is "
        "stopped at night. This confirms the safety-oriented behavior encoded in the "
        "state machine."
    ),
    (
        "The results demonstrate correct separation between signal conditioning and "
        "state-based decision making. The hysteresis structure prevents flicker, while "
        "Stateflow provides a readable and deterministic controller implementation."
    ),
]


REFERENCES = [
    "[1] MathWorks, \"Model-Based Design with Simulink,\" MathWorks Documentation. [Online]. Available: https://www.mathworks.com/help/simulink/gs/model-based-design.html",
    "[2] MathWorks, \"Stateflow Documentation,\" MathWorks Documentation. [Online]. Available: https://www.mathworks.com/help/stateflow/index.html",
    "[3] MathWorks, \"Design Finite State Machines in Stateflow,\" MathWorks Documentation. [Online]. Available: https://www.mathworks.com/help/stateflow/gs/get-started-introduction.html",
    "[4] K. M. F. Shahriar, \"Automatic headlight controlling of vehicle using ambient light sensor based on phototransistor,\" ScienceOpen Preprints, 2021.",
    "[5] B. H. Kim et al., \"Design and hardware-in-the-loop simulation of an automatic headlight control system,\" IFAC Proceedings Volumes, vol. 39, no. 16, pp. 746-751, 2006.",
    "[6] Y. L. Chen et al., \"New method of automatic control for vehicle headlights,\" Optik, vol. 157, pp. 718-723, 2018.",
]


SCENARIO_ROWS = [
    ("1", "Day + Moving", "800", "60", "0"),
    ("2", "Night + Moving", "100", "60", "1"),
    ("3", "Tunnel Entry", "800 -> 100 -> 800", "60", "0 -> 1 -> 0"),
    ("4", "Night + Stopped", "100", "0", "1"),
]


INPUT_OUTPUT_ROWS = [
    ("AmbientLux", "Input", "Brightness around the vehicle"),
    ("VehicleSpeed", "Input", "Vehicle moving or stopped status"),
    ("ManualOverride", "Input", "Driver command to force lamps ON"),
    ("Fault", "Input", "Fault supervision input, currently held at 0"),
    ("headlamp_cmd", "Output", "Final lamp command: 0 = OFF, 1 = ON"),
]


LOGIC_ROWS = [
    ("lux_lt_300", "AmbientLux < 300", "Detect dark condition"),
    ("lux_gt_500", "AmbientLux > 500", "Detect bright condition"),
    ("Unit Delay", "Stores previous lux_status", "Provides hysteresis memory"),
    ("NOT", "Inverts lux_gt_500", "Keeps dark status active until brightness is high"),
    ("AND", "previous_lux_status AND NOT lux_gt_500", "Maintains previous dark state"),
    ("OR", "... OR lux_lt_300", "Generates final lux_status"),
    ("speed_gt_0", "VehicleSpeed > 0", "Detect moving condition"),
]


STATE_ROWS = [
    ("OFF", "Bright condition", "headlamp_cmd = 0"),
    ("ON", "Dark and moving", "headlamp_cmd = 1"),
    ("HOLD", "Dark and stopped", "headlamp_cmd = 1"),
    ("MANUAL", "Manual override active", "headlamp_cmd = 1"),
]


FIGURES = [
    ("figure_top_level.png", "Figure 1. Top-Level Simulink Architecture"),
    ("figure_input_subsystem.png", "Figure 2. Input Subsystem"),
    ("figure_logic_subsystem.png", "Figure 3. Logic Subsystem"),
    ("figure_stateflow.png", "Figure 4. Stateflow Controller"),
    ("figure_output_subsystem.png", "Figure 5. Output Subsystem"),
    ("figure_scope_scenario1.png", "Figure 6. Scenario 1 Output"),
    ("figure_scope_scenario2.png", "Figure 7. Scenario 2 Output"),
    ("figure_scope_scenario3.png", "Figure 8. Scenario 3 Output"),
    ("figure_scope_scenario4.png", "Figure 9. Scenario 4 Output"),
]


@dataclass
class FigureSpec:
    filename: str
    caption: str
    width_inches: float = 5.8


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run.add_text("Update field in Word to populate the table of contents.")
    run._r.append(fld_end)


def set_document_margins(document: Document, *, top=0.8, bottom=0.8, left=0.9, right=0.8) -> None:
    for section in document.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_two_columns(section) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), "2")
    cols_el.set(qn("w:space"), "720")
    if not cols:
        sect_pr.append(cols_el)


def set_paragraph_font(paragraph, size: float = 12, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold if bold else run.bold
        run.italic = italic if italic else run.italic


def create_styles(document: Document, body_size: int = 12) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(body_size)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.font.size = Pt(size)

    if "CaptionCustom" not in document.styles:
        style = document.styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(max(body_size - 1, 10))
        style.font.italic = True


def add_header_footer(document: Document, header_text: str) -> None:
    for section in document.sections:
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_p.text = header_text
        set_paragraph_font(header_p, size=10)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.clear()
        footer_p.add_run("Page ")
        add_page_field(footer_p)
        set_paragraph_font(footer_p, size=10)


def add_body_paragraph(document: Document, text: str, *, justify: bool = True, size: int = 12, after_pt: int = 6) -> None:
    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1, *, center: bool = False, before: int = 10, after: int = 6):
    p = document.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.bold = True
    run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))
    return p


def build_placeholder_image(path: Path, label: str, size=(1600, 900)) -> None:
    bg = (245, 247, 250)
    border = (112, 128, 144)
    accent = (30, 41, 59)
    text_color = (51, 65, 85)
    image = Image.new("RGB", size, bg)
    draw = ImageDraw.Draw(image)
    margin = 70
    draw.rounded_rectangle(
        [margin, margin, size[0] - margin, size[1] - margin],
        radius=24,
        outline=border,
        width=6,
        fill=(255, 255, 255),
    )
    draw.line(
        [(margin + 30, 180), (size[0] - margin - 30, 180)],
        fill=(203, 213, 225),
        width=3,
    )
    try:
        title_font = ImageFont.truetype("times.ttf", 56)
        body_font = ImageFont.truetype("times.ttf", 34)
    except OSError:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.text((margin + 40, 95), PROJECT_NAME, fill=accent, font=title_font)
    draw.text((margin + 40, 260), label, fill=text_color, font=title_font)
    note = "Replace this placeholder with the corresponding model screenshot or simulation result."
    draw.text((margin + 40, 360), note, fill=text_color, font=body_font)
    draw.text((margin + 40, 430), "Preferred image: clean Simulink screenshot or scope capture.", fill=text_color, font=body_font)
    image.save(path)


def create_all_placeholder_images() -> dict[str, Path]:
    mapping = {}
    for filename, caption in FIGURES:
        path = IMG_DIR / filename
        build_placeholder_image(path, caption)
        mapping[caption] = path
    return mapping


def add_figure(document: Document, image_path: Path, caption: str, *, width_inches: float = 5.7) -> None:
    p_img = document.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    p_cap = document.add_paragraph(style="CaptionCustom")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(8)
    p_cap.add_run(caption)
    set_paragraph_font(p_cap, size=11, italic=True)


def add_table(document: Document, headers: Iterable[str], rows: Iterable[Iterable[str]], *, col_widths=None) -> None:
    table = document.add_table(rows=1, cols=len(list(headers)))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], str(header), bold=True, center=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        header_cells[idx]._tc.get_or_add_tcPr().append(shading)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), center=idx == 0)
    table.autofit = True
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    document.add_paragraph()


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(INST_PLACEHOLDER)
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.bold = True

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(28)
    r2 = p2.add_run(TITLE_TEXT)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2.font.size = Pt(22)
    r2.bold = True

    meta = [
        f"Course: {COURSE_TEXT}",
        f"Department: {DEPT_PLACEHOLDER}",
        f"Submitted By: {AUTHOR_PLACEHOLDER}",
        f"PRN: {PRN_PLACEHOLDER}",
    ]
    for line in meta:
        p_line = document.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_before = Pt(12)
        r = p_line.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(14)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(110)
    r3 = p3.add_run("Project Report")
    r3.font.name = "Times New Roman"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r3.font.size = Pt(18)
    r3.bold = True

    document.add_page_break()


def create_report_doc(figures: dict[str, Path]) -> None:
    document = Document()
    ensure_dirs()
    set_document_margins(document)
    create_styles(document, body_size=12)
    add_header_footer(document, f"{PROJECT_NAME} - Project Report")

    add_cover_page(document)

    add_heading(document, "Table of Contents", level=1)
    add_toc_field(document.add_paragraph())
    document.add_page_break()

    add_heading(document, "Abstract", level=1)
    add_body_paragraph(document, ABSTRACT_TEXT)

    add_heading(document, "Introduction", level=1)
    for para in INTRO_PARAS:
        add_body_paragraph(document, para)

    add_heading(document, "Problem Statement", level=1)
    add_body_paragraph(document, PROBLEM_STATEMENT)

    add_heading(document, "System Architecture", level=1)
    add_body_paragraph(document, SYSTEM_ARCHITECTURE_TEXT)
    add_figure(document, figures["Figure 1. Top-Level Simulink Architecture"], "Figure 1: Top-Level Simulink Architecture")

    add_heading(document, "Inputs and Outputs", level=1)
    add_table(document, ["Signal", "Type", "Description"], INPUT_OUTPUT_ROWS, col_widths=[1.5, 1.0, 3.7])

    add_heading(document, "Input Subsystem", level=1)
    for para in INPUT_SUBSYSTEM_PARAS:
        add_body_paragraph(document, para)
    add_table(
        document,
        ["Scenario", "Driving Condition", "AmbientLux", "VehicleSpeed", "Expected headlamp_cmd"],
        SCENARIO_ROWS,
        col_widths=[0.7, 1.9, 1.5, 1.1, 1.7],
    )
    add_figure(document, figures["Figure 2. Input Subsystem"], "Figure 2: Input Subsystem")

    add_heading(document, "Logic Subsystem", level=1)
    for para in LOGIC_SUBSYSTEM_PARAS:
        add_body_paragraph(document, para)
    add_table(document, ["Block / Signal", "Logic", "Purpose"], LOGIC_ROWS, col_widths=[1.5, 2.3, 2.4])
    add_figure(document, figures["Figure 3. Logic Subsystem"], "Figure 3: Logic Subsystem")

    add_heading(document, "Stateflow Controller", level=1)
    for para in STATEFLOW_PARAS:
        add_body_paragraph(document, para)
    add_table(document, ["State", "Active Condition", "Output"], STATE_ROWS, col_widths=[1.0, 3.4, 1.7])
    add_figure(document, figures["Figure 4. Stateflow Controller"], "Figure 4: Stateflow Controller")

    add_heading(document, "Output Subsystem", level=1)
    for para in OUTPUT_PARAS:
        add_body_paragraph(document, para)
    add_figure(document, figures["Figure 5. Output Subsystem"], "Figure 5: Output Subsystem")

    add_heading(document, "Test Cases and Results", level=1)
    add_body_paragraph(
        document,
        "The model was simulated for four representative scenarios. The resulting headlamp behavior matched the intended controller logic in every case.",
    )
    add_table(
        document,
        ["Scenario", "Condition", "Expected Output"],
        [(row[0], row[1], row[4]) for row in SCENARIO_ROWS],
        col_widths=[0.9, 3.5, 1.6],
    )
    add_figure(document, figures["Figure 6. Scenario 1 Output"], "Figure 6: Scenario 1 Scope Output")
    add_figure(document, figures["Figure 7. Scenario 2 Output"], "Figure 7: Scenario 2 Scope Output")
    add_figure(document, figures["Figure 8. Scenario 3 Output"], "Figure 8: Scenario 3 Scope Output")
    add_figure(document, figures["Figure 9. Scenario 4 Output"], "Figure 9: Scenario 4 Scope Output")

    add_heading(document, "Conclusion", level=1)
    add_body_paragraph(document, CONCLUSION_TEXT)

    document.save(str(REPORT_FILE))


def add_ieee_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(TITLE_TEXT)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)
    run.bold = True

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(8)
    author.add_run(
        f"{AUTHOR_PLACEHOLDER}, {DEPT_PLACEHOLDER}, {INST_PLACEHOLDER}, "
        f"{COURSE_TEXT}, PRN: {PRN_PLACEHOLDER}"
    )
    set_paragraph_font(author, size=10)


def add_ieee_section(document: Document, heading: str, paragraphs: Iterable[str]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(heading)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = True

    for para in paragraphs:
        p_body = document.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_body.paragraph_format.space_after = Pt(4)
        run = p_body.add_run(para)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)


def add_ieee_figure(document: Document, image_path: Path, caption: str) -> None:
    p_img = document.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(3.15))

    p_cap = document.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(4)
    cap_run = p_cap.add_run(caption)
    cap_run.font.name = "Times New Roman"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap_run.font.size = Pt(9)
    cap_run.italic = True


def add_ieee_table(document: Document, title: str, headers: Iterable[str], rows: Iterable[Iterable[str]]) -> None:
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(9)
    title_run.bold = True

    add_table(document, headers, rows)


def add_ieee_compact_table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(title)
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(8.5)
    title_run.bold = True

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, center=True)
        for run in hdr[i].paragraphs[0].runs:
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, center=i == 0)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
    document.add_paragraph()


def add_ieee_placeholder(document: Document, caption: str) -> None:
    box = document.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    cell.width = Inches(2.9)
    set_cell_text(cell, f"[Insert {caption} Screenshot Here]", center=True)
    for run in cell.paragraphs[0].runs:
        run.font.size = Pt(8.5)
        run.italic = True
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(8.5)
    run.italic = True


def add_borderless_body_table(document: Document):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                edge = OxmlElement(f"w:{side}")
                edge.set(qn("w:val"), "nil")
                tc_borders.append(edge)
            tc_pr.append(tc_borders)
    return table


def add_cell_heading(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = True


def add_cell_paragraph(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9.5)


def create_ieee_doc(figures: dict[str, Path]) -> None:
    document = Document()
    ensure_dirs()
    set_document_margins(document, top=0.65, bottom=0.65, left=0.6, right=0.6)
    create_styles(document, body_size=10)
    add_header_footer(document, f"{PROJECT_NAME} - IEEE Paper")

    add_ieee_title_block(document)

    abstract_heading = document.add_paragraph()
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = abstract_heading.add_run("Abstract")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.bold = True

    abstract_p = document.add_paragraph()
    abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    abs_run = abstract_p.add_run(ABSTRACT_TEXT)
    abs_run.font.name = "Times New Roman"
    abs_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    abs_run.font.size = Pt(10)

    keywords = document.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.space_after = Pt(4)
    k1 = keywords.add_run("Keywords: ")
    k1.bold = True
    k1.font.name = "Times New Roman"
    k1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    k1.font.size = Pt(10)
    k2 = keywords.add_run(
        "Automatic headlight control, Simulink, Stateflow, automotive control, hysteresis, model-based design."
    )
    k2.font.name = "Times New Roman"
    k2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    k2.font.size = Pt(10)

    body_page_1 = add_borderless_body_table(document)
    left_1, right_1 = body_page_1.rows[0].cells
    left_1.paragraphs[0].clear()
    right_1.paragraphs[0].clear()

    add_cell_heading(left_1, "I. Introduction")
    for para in IEEE_INTRO:
        add_cell_paragraph(left_1, para)
    add_cell_heading(left_1, "II. Literature Review")
    for para in IEEE_LIT_REVIEW:
        add_cell_paragraph(left_1, para)

    add_cell_heading(right_1, "III. Methodology")
    for para in IEEE_METHOD:
        add_cell_paragraph(right_1, para)

    document.add_page_break()

    body_page_2 = add_borderless_body_table(document)
    left_2, right_2 = body_page_2.rows[0].cells
    left_2.paragraphs[0].clear()
    right_2.paragraphs[0].clear()

    add_cell_heading(left_2, "IV. Results and Discussion")
    for para in IEEE_RESULTS:
        add_cell_paragraph(left_2, para)
    add_cell_heading(left_2, "V. Conclusion")
    add_cell_paragraph(left_2, CONCLUSION_TEXT)

    add_cell_heading(right_2, "References")
    for ref in REFERENCES:
        add_cell_paragraph(right_2, ref)

    document.add_page_break()

    add_ieee_compact_table(
        document,
        "TABLE I. SYSTEM INPUTS AND OUTPUTS",
        ["Signal", "Role", "Meaning"],
        [
            ["AmbientLux", "Input", "Light level (lux)"],
            ["VehicleSpeed", "Input", "Moving or stopped"],
            ["ManualOverride", "Input", "Driver forces lamp ON"],
            ["Fault", "Input", "Supervisory fault input"],
            ["headlamp_cmd", "Output", "0 = OFF, 1 = ON"],
        ],
    )
    add_ieee_compact_table(
        document,
        "TABLE II. SCENARIO DEFINITIONS",
        ["Scenario", "Condition", "Output"],
        [
            ["1", "Day + Moving", "0"],
            ["2", "Night + Moving", "1"],
            ["3", "Tunnel Entry", "0 -> 1 -> 0"],
            ["4", "Night + Stopped", "1"],
        ],
    )
    add_ieee_placeholder(document, "Figure 1. Input Subsystem")
    add_ieee_placeholder(document, "Figure 2. Logic Subsystem")
    add_ieee_placeholder(document, "Figure 3. Stateflow Controller")
    add_ieee_placeholder(document, "Figure 4. Output Subsystem")
    add_ieee_placeholder(document, "Figure 5. Tunnel Entry Simulation Result")

    document.save(str(IEEE_FILE))


def main() -> None:
    ensure_dirs()
    figures = create_all_placeholder_images()
    create_report_doc(figures)
    create_ieee_doc(figures)
    print(f"Created: {REPORT_FILE}")
    print(f"Created: {IEEE_FILE}")


if __name__ == "__main__":
    main()
